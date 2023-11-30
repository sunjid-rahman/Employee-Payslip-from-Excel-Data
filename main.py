import pandas as pd
import datetime as dt
import os
from docx import Document
from docx2pdf import convert


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def convert_excel_to_doc_and_pdf():
    data_file = 'employee_data.xlsx'
    df = pd.read_excel(data_file)
    output_directory = str(dt.date.today())
    if not os.path.isdir(output_directory):
        os.mkdir(output_directory)
    output_directory_doc = output_directory + '/doc'
    output_directory_pdf = output_directory + '/pdf'
    if not os.path.isdir(output_directory_doc):
        os.mkdir(output_directory_doc)
    if not os.path.isdir(output_directory_pdf):
        os.mkdir(output_directory_pdf)

    for item, data in df.iterrows():
        template_file_path = 'payslip-template.docx'
        output_file_path = output_directory_doc + '/' + str(data['ID']) + '.docx'
        variables = {
            "${MONTH}": str(data['Month']),
            "${YEAR}": str(data['Year']),
            "${DATE}": str(data['Date']),
            "${NAME}": str(data['Name']),
            "${ID}": str(data['ID']),
            "${DESIGNATION}": data['Designation'],
            "${BASIC}": str(data['Basic']),
            "${HOUSERENT}": str(data[r'House Rent']),
            "${MEDICAL}": str(data['Medical']),
            "${CONVEYANCE}": str(data['Conveyance']),
            "${LFA}": str(data['LFA']),
            "${PAYMENTSUB}": str(data[r'Payment Subtotal']),
            "${PFCONTRIBUTION}": str(data[r'PF Contribution']),
            "${TAX}": str(data['Tax']),
            "${DEDUCTIONSUB}": str(data[r'Deduction Subtotal']),
            "${DEDUCTIONTOTAL}": str(data[r'Deduction Total']),
            "${TOTAL}": str(data['Total'])
        }
        template_document = Document(template_file_path)
        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        try:
            input_file = output_file_path
            output_file = output_directory_pdf + '/' + str(data['ID']) + '.pdf'
            convert(input_file, output_file)
        except Exception as e:
            print(e)


if __name__ == '__main__':
    convert_excel_to_doc_and_pdf()
